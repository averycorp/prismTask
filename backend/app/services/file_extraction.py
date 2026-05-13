"""File-content -> structured task extraction.

Companion to ``extract_tasks_from_text``, but accepts a binary file (text,
PDF, DOCX, XLSX, image) and asks Claude Haiku for a richer single-task
suggestion: title, description, due date, priority, project hint, tags,
detected dates, and subtasks.

The Android client wires this into the "Extract from file" affordance on
the task editor — drop in a screenshot, JSX, syllabus, or meeting notes
and the editor pre-fills as much as possible from the contents.

The companion router lives at ``POST /api/v1/ai/files/extract`` and the
response shape is ``app.schemas.ai.FileExtractionResponse``.
"""

from __future__ import annotations

import base64
import io
import json
import logging
from typing import Any, Optional

from app.services.ai_productivity import _get_client, _parse_ai_json, get_model

logger = logging.getLogger(__name__)

# Hard cap on the binary we accept (the router enforces this too — keeping
# it here makes the service safe to call directly from tests/scripts).
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB

# How much extracted text we forward to Claude. Haiku has a 200k window but
# we keep it well under that to bound token cost on a 10 MB DOCX/XLSX.
MAX_TEXT_CHARS = 50_000

TEXT_LIKE_MIMES = {
    "text/plain",
    "text/markdown",
    "text/x-markdown",
    "text/html",
    "text/css",
    "text/csv",
    "text/x-python",
    "text/javascript",
    "text/typescript",
    "text/x-kotlin",
    "text/x-java",
    "text/x-c",
    "text/x-c++",
    "text/x-go",
    "text/x-rust",
    "text/x-ruby",
    "text/x-shellscript",
    "text/x-yaml",
    "text/yaml",
    "application/json",
    "application/xml",
    "text/xml",
    "application/javascript",
    "application/typescript",
    "application/x-yaml",
    "application/x-sh",
    "application/x-httpd-php",
}

IMAGE_MIMES = {
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/gif",
    "image/webp",
}

PDF_MIMES = {"application/pdf"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
XLSX_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

# Browsers occasionally serve source files as application/octet-stream;
# the extension is the tiebreaker.
TEXT_EXTENSIONS: tuple[str, ...] = (
    ".jsx", ".tsx", ".ts", ".js", ".mjs", ".cjs",
    ".py", ".rb", ".go", ".rs", ".java", ".kt", ".kts", ".scala",
    ".c", ".cpp", ".cc", ".h", ".hpp", ".cs", ".swift", ".m", ".mm",
    ".sh", ".bash", ".zsh", ".ps1", ".bat",
    ".md", ".markdown", ".txt", ".rst", ".adoc",
    ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".env",
    ".html", ".htm", ".css", ".scss", ".sass", ".less",
    ".xml", ".svg", ".sql", ".graphql", ".proto",
    ".lua", ".dart", ".vue", ".astro",
)


def extract_from_file(
    file_bytes: bytes,
    filename: str,
    mime_type: str,
    tier: str = "FREE",
) -> dict[str, Any]:
    """Extract a structured task suggestion from a file upload.

    Routes the bytes to a text extractor (PDF / DOCX / XLSX / text-like)
    or to Claude's multimodal endpoint (images), then asks Haiku for a
    single-task JSON response. ``tier`` is accepted for symmetry with the
    other ai_productivity helpers; model selection is handled by
    ``get_model("file_extraction")`` and is currently Haiku for everyone.

    The response always carries a ``technical_metadata`` block populated
    deterministically from the file bytes (page counts, EXIF, doc-info,
    etc.) regardless of whether the LLM call succeeded — empty/oversize
    short-circuits still include what was readable before the bailout.
    """
    prompt_filename = (filename or "uploaded-file").strip() or "uploaded-file"
    mime = (mime_type or "application/octet-stream").lower()

    if not file_bytes:
        empty = _empty_response(prompt_filename, mime)
        empty["technical_metadata"] = _empty_technical_metadata(0)
        return empty
    if len(file_bytes) > MAX_FILE_BYTES:
        raise ValueError(
            f"File too large: {len(file_bytes)} bytes (max {MAX_FILE_BYTES})"
        )

    technical_metadata = _collect_technical_metadata(file_bytes, prompt_filename, mime)

    client = _get_client()
    model = get_model("file_extraction")

    if mime in IMAGE_MIMES:
        user_content = _build_image_message(file_bytes, prompt_filename, mime)
    else:
        body_text = _resolve_text_body(file_bytes, prompt_filename, mime)
        if not body_text.strip():
            empty = _empty_response(prompt_filename, mime)
            empty["technical_metadata"] = technical_metadata
            return empty
        if len(body_text) > MAX_TEXT_CHARS:
            body_text = body_text[:MAX_TEXT_CHARS] + "\n... [truncated]"
        # Refresh the text-like counts from the *parsed* body rather than the
        # raw bytes so PDF/DOCX/XLSX get meaningful word/line numbers.
        technical_metadata.update(_text_counts_for(body_text))
        user_content = [
            {"type": "text", "text": _build_prompt(prompt_filename, mime, body_text)}
        ]

    last_error: Exception | None = None
    for attempt in range(2):
        try:
            message = client.messages.create(
                model=model,
                max_tokens=2048,
                messages=[{"role": "user", "content": user_content}],
            )
            content = message.content[0].text
            parsed = _parse_ai_json(content)
            if not isinstance(parsed, dict):
                raise ValueError("Expected a JSON object")
            parsed.setdefault("source_file_name", prompt_filename)
            parsed.setdefault("source_mime_type", mime)
            parsed["technical_metadata"] = technical_metadata
            return parsed
        except (json.JSONDecodeError, KeyError, TypeError, IndexError, ValueError) as e:
            last_error = e
            logger.error(
                "File extraction parse failed (attempt %d): %s", attempt + 1, e
            )
            if attempt == 0:
                continue
            raise ValueError(f"Failed to parse AI response after retry: {e}") from e
        except Exception as e:
            logger.error("File extraction AI error: %s: %s", type(e).__name__, e)
            raise
    raise ValueError(f"Failed to parse AI response: {last_error}")


def _resolve_text_body(file_bytes: bytes, filename: str, mime: str) -> str:
    """Decode/parse file bytes to plain text. Empty string for unsupported types."""
    lower_name = filename.lower()
    if mime in PDF_MIMES or lower_name.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    if mime in DOCX_MIMES or lower_name.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)
    if mime in XLSX_MIMES or lower_name.endswith(".xlsx"):
        return _extract_text_from_xlsx(file_bytes)
    if _is_text_like(mime, filename):
        return file_bytes.decode("utf-8", errors="replace")
    # Last-ditch: many "application/octet-stream" uploads are actually text.
    try:
        decoded = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return ""
    return decoded if decoded.isprintable() or "\n" in decoded else ""


def _is_text_like(mime: str, filename: str) -> bool:
    if mime in TEXT_LIKE_MIMES or mime.startswith("text/"):
        return True
    return filename.lower().endswith(TEXT_EXTENSIONS)


def _extract_text_from_pdf(data: bytes) -> str:
    try:
        import pypdf
    except ImportError as e:
        raise RuntimeError("pypdf is not installed") from e
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001 — pypdf raises a wide tree
        logger.warning("PDF could not be opened: %s", e)
        return ""
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001
            logger.debug("PDF page extraction failed: %s", e)
    return "\n\n".join(p for p in parts if p).strip()


def _extract_text_from_docx(data: bytes) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed (add `python-docx` to requirements.txt)"
        ) from e
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001 — python-docx raises PackageNotFoundError + others
        logger.warning("DOCX could not be opened: %s", e)
        return ""
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_text_from_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        raise RuntimeError(
            "openpyxl is not installed (add `openpyxl` to requirements.txt)"
        ) from e
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as e:  # noqa: BLE001 — openpyxl raises InvalidFileException + others
        logger.warning("XLSX could not be opened: %s", e)
        return ""
    try:
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    finally:
        wb.close()


def _build_image_message(file_bytes: bytes, filename: str, mime: str) -> list[dict]:
    """Multimodal payload — Claude reads the image and the prompt together."""
    encoded = base64.standard_b64encode(file_bytes).decode("ascii")
    # image/jpg is not a real IANA type but browsers occasionally send it;
    # Anthropic only accepts image/jpeg, so normalize.
    media_type = "image/jpeg" if mime == "image/jpg" else mime
    return [
        {
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": media_type,
                "data": encoded,
            },
        },
        {
            "type": "text",
            "text": _build_prompt(
                filename,
                media_type,
                "[image contents — see attached image]",
            ),
        },
    ]


def _build_prompt(filename: str, mime: str, body_text: str) -> str:
    return f"""You are an assistant that turns an uploaded file into a single
structured task with optional subtasks. The user wants as much auto-fill as
possible from the file contents — title, description, due date, priority,
project hint, tags, subtasks the file implies, *and* rich enrichment
metadata (life category, estimated duration, location, recurrence hint,
reminder offset, links, contacts, key entities, document type, language).

File metadata:
- name: {filename}
- mime: {mime}

File contents:
---
{body_text}
---

Return ONLY valid JSON with this exact shape (no commentary, no markdown fences):

{{
  "title": "Concise imperative title in Title Case, under 12 words",
  "description": "Optional short summary (1-2 sentences) or null",
  "suggested_due_date": "YYYY-MM-DD or null",
  "suggested_priority": 0,
  "suggested_project": "Short project name or null",
  "tags": ["tag1", "tag2"],
  "subtasks": [
    {{"title": "Subtask in imperative Title Case", "suggested_due_date": null}}
  ],
  "detected_dates": ["YYYY-MM-DD"],
  "confidence": 0.0,
  "notes": "Optional 1-sentence rationale or null",

  "life_category": "WORK | PERSONAL | SELF_CARE | HEALTH | UNCATEGORIZED",
  "estimated_duration_minutes": null,
  "recurrence_hint": "Free-text recurrence like 'every Monday', 'weekly', or null",
  "location": "Physical address / room / venue / video-call URL, or null",
  "reminder_offset_minutes": null,
  "urls": ["https://..."],
  "contacts": [
    {{"name": "Jane Doe or null", "email": "x@y.com or null", "phone": "+1... or null"}}
  ],
  "key_entities": ["Org or person or product names mentioned"],
  "document_type": "syllabus | meeting_notes | invoice | shopping_list | code_todos | report | event_invite | personal_note | recipe | article | other",
  "action_or_info": "action | info",
  "language": "ISO 639-1 code (en, es, fr, ...) or null"
}}

Rules:
- ``suggested_priority`` is 0 (none), 1 (low), 2 (medium), 3 (high), 4 (urgent).
- Pull every date the file references (deadlines, meeting dates, due dates,
  EXIF-style timestamps embedded in filenames). List all of them in
  ``detected_dates`` and pick the soonest forward-looking one as
  ``suggested_due_date``. If only past dates are present,
  ``suggested_due_date`` is null.
- Subtasks should reflect actionable lines in the file: TODO/FIXME comments
  in code, bullet items in notes, headings in syllabi, line items in
  spreadsheets. Cap at 20 subtasks.
- Tags are lowercase, no leading ``#``.
- ``life_category`` reflects the dominant life area the task belongs to.
  WORK = paid/professional work, code, meetings; PERSONAL = errands, family,
  social plans; SELF_CARE = hobbies, leisure, rest, journaling;
  HEALTH = exercise, doctor visits, medication, sleep, nutrition;
  UNCATEGORIZED only if genuinely ambiguous.
- ``estimated_duration_minutes`` is the realistic time to do the task end
  to end, between 0 and 1440. Null if the file gives no useful signal.
- ``recurrence_hint`` is the *literal phrasing* of any recurrence the file
  implies ("every Monday", "weekly stand-up", "biweekly retro"). The user
  confirms before any RecurrenceRule is built — don't fabricate.
- ``location`` may be an address, room number, venue name, or a meeting
  URL (Zoom/Meet/etc). Null if nothing location-shaped is in the file.
- ``reminder_offset_minutes`` is set only when the file mentions a
  specific alert ahead of time ("remind me 15 minutes before"). Null
  otherwise.
- ``urls`` is the deduped list of links found in the file body.
- ``contacts`` extracts people referenced by name + email/phone. Skip
  contact entries with only a name and no contact info.
- ``key_entities`` is a short list (<=10) of named entities — people,
  organizations, product or project names — that ground the task.
- ``document_type`` classifies what kind of document this is. Use
  ``other`` only when nothing else fits.
- ``action_or_info`` is ``action`` when the document implies the user
  needs to do something, ``info`` when it's reference material.
- ``language`` is the dominant language of the body in ISO 639-1.
- If the file is empty or contains nothing actionable, return ``title=""``
  and ``confidence=0.0`` — never invent content. Enrichment fields are
  null/empty in that case.
"""


def _empty_response(filename: str, mime: str) -> dict[str, Any]:
    return {
        "title": "",
        "description": None,
        "suggested_due_date": None,
        "suggested_priority": 0,
        "suggested_project": None,
        "tags": [],
        "subtasks": [],
        "detected_dates": [],
        "confidence": 0.0,
        "notes": None,
        "source_file_name": filename,
        "source_mime_type": mime,
        "life_category": None,
        "estimated_duration_minutes": None,
        "recurrence_hint": None,
        "location": None,
        "reminder_offset_minutes": None,
        "urls": [],
        "contacts": [],
        "key_entities": [],
        "document_type": None,
        "action_or_info": None,
        "language": None,
        "technical_metadata": None,
    }


# --- Deterministic technical metadata --------------------------------------


def _empty_technical_metadata(file_size_bytes: int) -> dict[str, Any]:
    """Skeleton dict matching ``FileTechnicalMetadata`` with everything null/empty."""
    return {
        "file_size_bytes": file_size_bytes,
        "page_count": None,
        "doc_title": None,
        "doc_author": None,
        "doc_subject": None,
        "doc_keywords": None,
        "doc_creation_date": None,
        "doc_modification_date": None,
        "doc_last_modified_by": None,
        "doc_revision": None,
        "paragraph_count": None,
        "table_count": None,
        "sheet_names": [],
        "sheet_count": None,
        "row_count_total": None,
        "width_px": None,
        "height_px": None,
        "image_taken_at": None,
        "camera_make": None,
        "camera_model": None,
        "gps_lat": None,
        "gps_lon": None,
        "line_count": None,
        "word_count": None,
        "char_count": None,
    }


def _collect_technical_metadata(
    file_bytes: bytes, filename: str, mime: str
) -> dict[str, Any]:
    """Dispatch to a per-type collector and return a ``FileTechnicalMetadata``-
    shaped dict. Best-effort: any per-collector failure logs and is skipped
    so the caller still gets the parts that succeeded."""
    meta = _empty_technical_metadata(len(file_bytes))
    lower_name = filename.lower()
    try:
        if mime in PDF_MIMES or lower_name.endswith(".pdf"):
            _collect_pdf_metadata(file_bytes, meta)
        elif mime in DOCX_MIMES or lower_name.endswith(".docx"):
            _collect_docx_metadata(file_bytes, meta)
        elif mime in XLSX_MIMES or lower_name.endswith(".xlsx"):
            _collect_xlsx_metadata(file_bytes, meta)
        elif mime in IMAGE_MIMES:
            _collect_image_metadata(file_bytes, meta)
        elif _is_text_like(mime, filename):
            # Word/line/char counts for raw text — re-computed on the parsed
            # body later for PDF/DOCX/XLSX inside ``extract_from_file``.
            meta.update(_text_counts_for(file_bytes.decode("utf-8", errors="replace")))
    except Exception as e:  # noqa: BLE001 — never let metadata break extraction
        logger.warning("Technical metadata collection failed for %s: %s", filename, e)
    return meta


def _text_counts_for(body_text: str) -> dict[str, int]:
    return {
        "char_count": len(body_text),
        "line_count": body_text.count("\n") + (1 if body_text and not body_text.endswith("\n") else 0),
        "word_count": len(body_text.split()),
    }


def _collect_pdf_metadata(data: bytes, meta: dict[str, Any]) -> None:
    try:
        import pypdf
    except ImportError:
        return
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        logger.debug("PDF metadata: open failed: %s", e)
        return
    try:
        meta["page_count"] = len(reader.pages)
    except Exception as e:  # noqa: BLE001
        logger.debug("PDF metadata: page count failed: %s", e)
    try:
        info = reader.metadata
    except Exception as e:  # noqa: BLE001
        logger.debug("PDF metadata: doc info failed: %s", e)
        info = None
    if info:
        meta["doc_title"] = _coerce_str(getattr(info, "title", None))
        meta["doc_author"] = _coerce_str(getattr(info, "author", None))
        meta["doc_subject"] = _coerce_str(getattr(info, "subject", None))
        # ``keywords`` and creation/mod dates come back as strings or
        # IndirectObject wrappers — coerce both.
        keywords = info.get("/Keywords") if hasattr(info, "get") else None
        meta["doc_keywords"] = _coerce_str(keywords)
        meta["doc_creation_date"] = _coerce_str(
            info.get("/CreationDate") if hasattr(info, "get") else None
        )
        meta["doc_modification_date"] = _coerce_str(
            info.get("/ModDate") if hasattr(info, "get") else None
        )


def _collect_docx_metadata(data: bytes, meta: dict[str, Any]) -> None:
    try:
        from docx import Document
    except ImportError:
        return
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        logger.debug("DOCX metadata: open failed: %s", e)
        return
    try:
        meta["paragraph_count"] = sum(1 for p in doc.paragraphs if p.text.strip())
        meta["table_count"] = len(doc.tables)
    except Exception as e:  # noqa: BLE001
        logger.debug("DOCX metadata: counts failed: %s", e)
    try:
        core = doc.core_properties
    except Exception as e:  # noqa: BLE001
        logger.debug("DOCX metadata: core props failed: %s", e)
        return
    meta["doc_title"] = _coerce_str(getattr(core, "title", None))
    meta["doc_author"] = _coerce_str(getattr(core, "author", None))
    meta["doc_subject"] = _coerce_str(getattr(core, "subject", None))
    meta["doc_keywords"] = _coerce_str(getattr(core, "keywords", None))
    meta["doc_last_modified_by"] = _coerce_str(getattr(core, "last_modified_by", None))
    revision = getattr(core, "revision", None)
    if isinstance(revision, int) and revision >= 0:
        meta["doc_revision"] = revision
    meta["doc_creation_date"] = _coerce_datetime(getattr(core, "created", None))
    meta["doc_modification_date"] = _coerce_datetime(getattr(core, "modified", None))


def _collect_xlsx_metadata(data: bytes, meta: dict[str, Any]) -> None:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return
    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as e:  # noqa: BLE001
        logger.debug("XLSX metadata: open failed: %s", e)
        return
    try:
        names = [s.title for s in wb.worksheets]
        meta["sheet_names"] = names
        meta["sheet_count"] = len(names)
        total_rows = 0
        for sheet in wb.worksheets:
            mr = sheet.max_row
            if isinstance(mr, int):
                total_rows += mr
        meta["row_count_total"] = total_rows
        try:
            props = wb.properties
        except Exception:  # noqa: BLE001
            props = None
        if props is not None:
            meta["doc_title"] = _coerce_str(getattr(props, "title", None))
            meta["doc_author"] = _coerce_str(getattr(props, "creator", None))
            meta["doc_subject"] = _coerce_str(getattr(props, "subject", None))
            meta["doc_keywords"] = _coerce_str(getattr(props, "keywords", None))
            meta["doc_last_modified_by"] = _coerce_str(
                getattr(props, "lastModifiedBy", None)
            )
            meta["doc_creation_date"] = _coerce_datetime(getattr(props, "created", None))
            meta["doc_modification_date"] = _coerce_datetime(getattr(props, "modified", None))
    except Exception as e:  # noqa: BLE001
        logger.debug("XLSX metadata: enumeration failed: %s", e)
    finally:
        try:
            wb.close()
        except Exception:  # noqa: BLE001
            pass


def _collect_image_metadata(data: bytes, meta: dict[str, Any]) -> None:
    try:
        from PIL import ExifTags, Image
    except ImportError:
        return
    try:
        img = Image.open(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        logger.debug("Image metadata: open failed: %s", e)
        return
    try:
        meta["width_px"], meta["height_px"] = img.size
    except Exception:  # noqa: BLE001
        pass
    exif_raw: dict | None = None
    try:
        # ``_getexif`` is the legacy attribute on JPEG; ``.getexif()`` works
        # on most modern Pillow versions but returns an empty mapping for
        # non-EXIF formats (PNG/GIF/WEBP usually).
        exif_raw = img.getexif()
    except Exception as e:  # noqa: BLE001
        logger.debug("Image metadata: exif read failed: %s", e)
        exif_raw = None
    if not exif_raw:
        return
    tag_map = {v: k for k, v in ExifTags.TAGS.items()}  # name -> tag id
    gps_tag_id = tag_map.get("GPSInfo")
    for tag_id, value in exif_raw.items():
        name = ExifTags.TAGS.get(tag_id, str(tag_id))
        if name == "DateTimeOriginal" or name == "DateTime":
            meta.setdefault("image_taken_at", _coerce_str(value))
        elif name == "Make":
            meta["camera_make"] = _coerce_str(value)
        elif name == "Model":
            meta["camera_model"] = _coerce_str(value)
    if gps_tag_id is not None:
        gps_info = exif_raw.get(gps_tag_id)
        if gps_info:
            try:
                lat, lon = _parse_gps(gps_info)
                if lat is not None:
                    meta["gps_lat"] = lat
                if lon is not None:
                    meta["gps_lon"] = lon
            except Exception as e:  # noqa: BLE001
                logger.debug("Image metadata: gps parse failed: %s", e)


def _parse_gps(gps_info: Any) -> tuple[Optional[float], Optional[float]]:
    """Parse an EXIF GPSInfo IFD into decimal lat/lon (signed)."""
    from PIL import ExifTags

    gps_name_map = ExifTags.GPSTAGS
    fields: dict[str, Any] = {}
    # ``gps_info`` may be a dict-like already or an IFD object — both expose
    # ``.items()``. Build a name-keyed dict to read uniformly.
    try:
        items = gps_info.items()
    except AttributeError:
        return (None, None)
    for k, v in items:
        fields[gps_name_map.get(k, str(k))] = v

    def _dms_to_decimal(dms: Any) -> Optional[float]:
        try:
            d, m, s = dms
            return float(d) + float(m) / 60.0 + float(s) / 3600.0
        except Exception:  # noqa: BLE001
            return None

    lat = _dms_to_decimal(fields.get("GPSLatitude"))
    lat_ref = fields.get("GPSLatitudeRef")
    if lat is not None and isinstance(lat_ref, str) and lat_ref.upper() == "S":
        lat = -lat
    lon = _dms_to_decimal(fields.get("GPSLongitude"))
    lon_ref = fields.get("GPSLongitudeRef")
    if lon is not None and isinstance(lon_ref, str) and lon_ref.upper() == "W":
        lon = -lon
    return (lat, lon)


def _coerce_str(value: Any) -> Optional[str]:
    """Normalize a doc-info value to a clean string or None."""
    if value is None:
        return None
    if isinstance(value, bytes):
        try:
            value = value.decode("utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            return None
    text = str(value).strip()
    return text or None


def _coerce_datetime(value: Any) -> Optional[str]:
    """Coerce a ``datetime`` / string to ISO-8601 (or pass through strings)."""
    if value is None:
        return None
    try:
        from datetime import datetime
        if isinstance(value, datetime):
            return value.isoformat()
    except Exception:  # noqa: BLE001
        pass
    return _coerce_str(value)
