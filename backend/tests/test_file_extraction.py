"""Tests for the file -> task extraction service and router.

Service tests run against a stubbed Anthropic client. Router tests use
the shared httpx AsyncClient + auth_headers fixtures from conftest.
"""

import io
import json
import struct
import sys
import types
import zlib
from unittest.mock import MagicMock, patch

import pytest
from httpx import AsyncClient


def _png_bytes(width: int = 4, height: int = 3) -> bytes:
    """Minimal valid PNG so PIL can read width/height. No EXIF — that's
    exercised by image_metadata_extraction_handles_missing_exif."""
    sig = b"\x89PNG\r\n\x1a\n"

    def _chunk(tag: bytes, payload: bytes) -> bytes:
        crc = zlib.crc32(tag + payload) & 0xFFFFFFFF
        return struct.pack(">I", len(payload)) + tag + payload + struct.pack(">I", crc)

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b""
    for _ in range(height):
        raw += b"\x00" + b"\xff\x00\x00" * width
    idat = zlib.compress(raw)
    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


def _make_mock_message(payload: dict) -> MagicMock:
    """Shape a fake Anthropic Message whose first content block is the JSON
    payload our prompt asks the model to return.
    """
    block = MagicMock()
    block.text = json.dumps(payload)
    message = MagicMock()
    message.content = [block]
    return message


@pytest.fixture(autouse=True)
def mock_anthropic_module():
    mock_mod = types.ModuleType("anthropic")
    mock_mod.Anthropic = MagicMock  # type: ignore
    mock_mod.APIError = Exception  # type: ignore
    sys.modules["anthropic"] = mock_mod

    import importlib

    import app.services.ai_productivity
    importlib.reload(app.services.ai_productivity)
    import app.services.file_extraction  # noqa: F401 — re-import after reload
    importlib.reload(app.services.file_extraction)

    yield mock_mod

    if "anthropic" in sys.modules and sys.modules["anthropic"] is mock_mod:
        del sys.modules["anthropic"]
    importlib.reload(app.services.ai_productivity)
    importlib.reload(app.services.file_extraction)


class TestFileExtractionService:
    """Unit tests for ``extract_from_file``."""

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_text_file_routes_through_text_path(self):
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client
            mock_client.messages.create.return_value = _make_mock_message({
                "title": "Ship the Login Flow",
                "description": "Implement the login form and wire it up.",
                "suggested_due_date": "2026-05-20",
                "suggested_priority": 3,
                "suggested_project": "Auth",
                "tags": ["frontend", "auth"],
                "subtasks": [
                    {"title": "Wire Up Form Validation", "suggested_due_date": None},
                    {"title": "Add Forgot-Password Link", "suggested_due_date": None},
                ],
                "detected_dates": ["2026-05-20"],
                "confidence": 0.85,
                "notes": "Notes file describes a login feature.",
            })

            body = b"# Login feature\n\n- TODO wire validation\n- TODO add forgot link\n"
            result = extract_from_file(
                file_bytes=body,
                filename="login.md",
                mime_type="text/markdown",
            )

            assert result["title"] == "Ship the Login Flow"
            assert result["suggested_priority"] == 3
            assert len(result["subtasks"]) == 2
            assert result["source_file_name"] == "login.md"
            assert result["source_mime_type"] == "text/markdown"

            # The user-message should have routed through the text path —
            # one text block, no image block.
            (call,) = mock_client.messages.create.call_args_list
            messages = call.kwargs["messages"]
            assert len(messages) == 1
            blocks = messages[0]["content"]
            assert all(b["type"] == "text" for b in blocks)

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_image_file_routes_through_multimodal_path(self):
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client
            mock_client.messages.create.return_value = _make_mock_message({
                "title": "Follow Up On Meeting",
                "subtasks": [],
                "tags": [],
                "detected_dates": [],
                "confidence": 0.6,
            })

            # Tiny PNG-ish bytes; the mock doesn't actually parse them.
            result = extract_from_file(
                file_bytes=b"\x89PNG\r\n\x1a\n",
                filename="meeting.png",
                mime_type="image/png",
            )
            assert result["title"] == "Follow Up On Meeting"

            (call,) = mock_client.messages.create.call_args_list
            blocks = call.kwargs["messages"][0]["content"]
            types_seen = [b["type"] for b in blocks]
            assert "image" in types_seen
            assert "text" in types_seen

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_image_jpg_normalized_to_jpeg(self):
        """``image/jpg`` is not a real IANA type — Anthropic rejects it.
        The service must normalize to ``image/jpeg`` before sending."""
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client
            mock_client.messages.create.return_value = _make_mock_message({
                "title": "X", "subtasks": [], "tags": [],
                "detected_dates": [], "confidence": 0.1,
            })

            extract_from_file(
                file_bytes=b"\xff\xd8\xff\xe0",
                filename="photo.jpg",
                mime_type="image/jpg",
            )
            blocks = mock_client.messages.create.call_args.kwargs["messages"][0]["content"]
            image_block = next(b for b in blocks if b["type"] == "image")
            assert image_block["source"]["media_type"] == "image/jpeg"

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_jsx_extension_treated_as_text_even_with_octet_mime(self):
        """Browsers serve .jsx as application/octet-stream — extension wins."""
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client
            mock_client.messages.create.return_value = _make_mock_message({
                "title": "Refactor LoginForm",
                "subtasks": [{"title": "Extract Validation Helper"}],
                "tags": ["refactor"],
                "detected_dates": [], "confidence": 0.7,
            })

            jsx = b"// TODO: extract validation helper\nexport function LoginForm() { return null; }\n"
            result = extract_from_file(
                file_bytes=jsx,
                filename="LoginForm.jsx",
                mime_type="application/octet-stream",
            )
            assert result["title"] == "Refactor LoginForm"
            assert len(result["subtasks"]) == 1

    def test_empty_file_returns_empty_response_without_calling_claude(self):
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client

            result = extract_from_file(
                file_bytes=b"",
                filename="empty.txt",
                mime_type="text/plain",
            )
            assert result["title"] == ""
            assert result["confidence"] == 0.0
            assert result["source_file_name"] == "empty.txt"
            mock_client.messages.create.assert_not_called()

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_oversize_file_raises_value_error(self):
        from app.services.file_extraction import MAX_FILE_BYTES, extract_from_file

        oversized = b"x" * (MAX_FILE_BYTES + 1)
        with pytest.raises(ValueError, match="too large"):
            extract_from_file(
                file_bytes=oversized,
                filename="huge.txt",
                mime_type="text/plain",
            )

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_pdf_extraction_short_circuits_to_empty_when_no_text(self):
        """If pypdf yields no text we skip the Claude call entirely."""
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic, \
             patch("app.services.file_extraction._extract_text_from_pdf", return_value=""):
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client

            result = extract_from_file(
                file_bytes=b"%PDF-1.4 fake",
                filename="empty.pdf",
                mime_type="application/pdf",
            )
            assert result["title"] == ""
            mock_client.messages.create.assert_not_called()


class TestExtractFromFileRouter:
    """Integration tests for ``POST /api/v1/ai/files/extract``.

    The endpoint sits behind the AI feature gate (Pro-tier only), so these
    tests use ``pro_auth_headers`` instead of the bare ``auth_headers``
    fixture — auth_headers users get 403 from the AI middleware.
    """

    @pytest.mark.asyncio
    async def test_returns_structured_suggestion_for_text_file(
        self, client: AsyncClient, pro_auth_headers: dict
    ):
        with patch("app.services.file_extraction.extract_from_file") as mock_svc:
            mock_svc.return_value = {
                "title": "Plan Sprint",
                "description": "Outline the next two-week sprint.",
                "suggested_due_date": "2026-05-20",
                "suggested_priority": 2,
                "suggested_project": "Planning",
                "tags": ["sprint", "planning"],
                "subtasks": [
                    {"title": "Draft Goals", "suggested_due_date": None},
                    {"title": "Schedule Kickoff", "suggested_due_date": "2026-05-15"},
                ],
                "detected_dates": ["2026-05-15", "2026-05-20"],
                "confidence": 0.78,
                "notes": "Markdown notes file with TODOs",
                "source_file_name": "sprint.md",
                "source_mime_type": "text/markdown",
            }

            files = {"file": ("sprint.md", io.BytesIO(b"- TODO Draft goals"), "text/markdown")}
            resp = await client.post(
                "/api/v1/ai/files/extract",
                files=files,
                headers=pro_auth_headers,
            )
            assert resp.status_code == 200, resp.text
            body = resp.json()
            assert body["title"] == "Plan Sprint"
            assert body["suggested_priority"] == 2
            assert len(body["subtasks"]) == 2
            assert body["subtasks"][1]["suggested_due_date"] == "2026-05-15"
            assert body["detected_dates"] == ["2026-05-15", "2026-05-20"]
            assert body["source_file_name"] == "sprint.md"

    @pytest.mark.asyncio
    async def test_drops_subtasks_with_blank_titles(
        self, client: AsyncClient, pro_auth_headers: dict
    ):
        with patch("app.services.file_extraction.extract_from_file") as mock_svc:
            mock_svc.return_value = {
                "title": "X",
                "subtasks": [
                    {"title": "Real Subtask"},
                    {"title": "   "},
                    {"suggested_due_date": "2026-05-12"},  # missing title
                    "not a dict",
                ],
                "tags": ["#TaggedWithHash", "  ", "clean-tag"],
                "detected_dates": [],
                "confidence": 0.5,
            }

            files = {"file": ("x.txt", io.BytesIO(b"x"), "text/plain")}
            resp = await client.post(
                "/api/v1/ai/files/extract",
                files=files,
                headers=pro_auth_headers,
            )
            assert resp.status_code == 200, resp.text
            body = resp.json()
            assert [s["title"] for s in body["subtasks"]] == ["Real Subtask"]
            # Hash stripped; blank dropped.
            assert body["tags"] == ["TaggedWithHash", "clean-tag"]

    @pytest.mark.asyncio
    async def test_oversize_upload_returns_413(
        self, client: AsyncClient, pro_auth_headers: dict
    ):
        from app.routers.ai import FILE_EXTRACT_MAX_BYTES

        big = b"x" * (FILE_EXTRACT_MAX_BYTES + 1)
        files = {"file": ("big.txt", io.BytesIO(big), "text/plain")}
        resp = await client.post(
            "/api/v1/ai/files/extract",
            files=files,
            headers=pro_auth_headers,
        )
        assert resp.status_code == 413

    @pytest.mark.asyncio
    async def test_service_unavailable_returns_503(
        self, client: AsyncClient, pro_auth_headers: dict
    ):
        with patch(
            "app.services.file_extraction.extract_from_file",
            side_effect=RuntimeError("ANTHROPIC_API_KEY not set"),
        ):
            files = {"file": ("x.txt", io.BytesIO(b"x"), "text/plain")}
            resp = await client.post(
                "/api/v1/ai/files/extract",
                files=files,
                headers=pro_auth_headers,
            )
            assert resp.status_code == 503

    @pytest.mark.asyncio
    async def test_unparseable_ai_response_returns_422(
        self, client: AsyncClient, pro_auth_headers: dict
    ):
        with patch(
            "app.services.file_extraction.extract_from_file",
            side_effect=ValueError("Failed to parse AI response after retry"),
        ):
            files = {"file": ("x.txt", io.BytesIO(b"x"), "text/plain")}
            resp = await client.post(
                "/api/v1/ai/files/extract",
                files=files,
                headers=pro_auth_headers,
            )
            assert resp.status_code == 422

    @pytest.mark.asyncio
    async def test_unauthenticated_request_is_rejected(self, client: AsyncClient):
        files = {"file": ("x.txt", io.BytesIO(b"x"), "text/plain")}
        resp = await client.post("/api/v1/ai/files/extract", files=files)
        assert resp.status_code in (401, 403)


class TestTechnicalMetadataCollection:
    """Unit tests for the deterministic file-side metadata collector.

    These don't touch Anthropic — they exercise the per-type collectors
    directly and assert the shape of the returned dict.
    """

    def test_text_file_carries_counts(self):
        from app.services.file_extraction import _collect_technical_metadata

        body = b"line one\nline two\nline three\n"
        meta = _collect_technical_metadata(body, "notes.txt", "text/plain")
        assert meta["file_size_bytes"] == len(body)
        assert meta["line_count"] == 3
        assert meta["word_count"] == 6
        assert meta["char_count"] == len(body.decode("utf-8"))
        assert meta["page_count"] is None  # not a PDF

    def test_empty_dict_for_unknown_binary_type(self):
        """Unknown binary mime: collector still fills file_size_bytes but
        leaves everything else empty/null. No exception."""
        from app.services.file_extraction import _collect_technical_metadata

        meta = _collect_technical_metadata(
            b"\x00\x01\x02", "weird.bin", "application/x-weird"
        )
        assert meta["file_size_bytes"] == 3
        assert meta["page_count"] is None
        assert meta["word_count"] is None
        assert meta["sheet_names"] == []
        assert meta["width_px"] is None

    def test_png_collects_dimensions(self):
        try:
            from PIL import Image  # noqa: F401
        except ImportError:
            pytest.skip("Pillow not installed in this environment")
        from app.services.file_extraction import _collect_technical_metadata

        png = _png_bytes(width=4, height=3)
        meta = _collect_technical_metadata(png, "tiny.png", "image/png")
        assert meta["width_px"] == 4
        assert meta["height_px"] == 3
        # PNG without EXIF — those fields stay null.
        assert meta["camera_make"] is None
        assert meta["gps_lat"] is None

    def test_docx_collects_paragraphs_and_core_props(self):
        from app.services.file_extraction import _collect_technical_metadata

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Hello world")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("")  # blank
        doc.core_properties.author = "Test Author"
        doc.core_properties.title = "My Doc"

        buf = io.BytesIO()
        doc.save(buf)

        meta = _collect_technical_metadata(
            buf.getvalue(),
            "doc.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        # Two non-empty paragraphs.
        assert meta["paragraph_count"] == 2
        assert meta["doc_author"] == "Test Author"
        assert meta["doc_title"] == "My Doc"
        assert meta["table_count"] == 0

    def test_xlsx_collects_sheet_names(self):
        from app.services.file_extraction import _collect_technical_metadata

        try:
            from openpyxl import Workbook
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = Workbook()
        wb.active.title = "Summary"
        wb.create_sheet("Detail")
        wb.create_sheet("Notes")
        # Write a couple of rows so max_row > 0.
        wb["Summary"]["A1"] = "Header"
        wb["Summary"]["A2"] = "Value"

        buf = io.BytesIO()
        wb.save(buf)

        meta = _collect_technical_metadata(
            buf.getvalue(),
            "report.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        assert meta["sheet_names"] == ["Summary", "Detail", "Notes"]
        assert meta["sheet_count"] == 3
        assert (meta["row_count_total"] or 0) >= 2


class TestEnrichmentFieldsInResponse:
    """Service-level test that the LLM-returned enrichment fields make it
    back to the caller, and the technical_metadata block is always attached."""

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_enrichment_fields_round_trip(self):
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client
            mock_client.messages.create.return_value = _make_mock_message({
                "title": "Weekly Team Standup",
                "description": "Sync on cross-team blockers.",
                "suggested_due_date": "2026-05-20",
                "suggested_priority": 2,
                "suggested_project": "Standups",
                "tags": ["work", "meeting"],
                "subtasks": [],
                "detected_dates": ["2026-05-20"],
                "confidence": 0.9,
                "notes": "Recurring meeting notes",
                "life_category": "WORK",
                "estimated_duration_minutes": 30,
                "recurrence_hint": "every Tuesday",
                "location": "https://meet.example.com/abc",
                "reminder_offset_minutes": 15,
                "urls": ["https://meet.example.com/abc"],
                "contacts": [
                    {"name": "Alex Doe", "email": "alex@example.com", "phone": None}
                ],
                "key_entities": ["Platform Team"],
                "document_type": "meeting_notes",
                "action_or_info": "action",
                "language": "en",
            })

            body = (
                "Weekly team standup - Tuesdays 10am\n"
                "Link: https://meet.example.com/abc\n"
                "Alex: alex@example.com\n"
            ).encode("utf-8")
            result = extract_from_file(
                file_bytes=body,
                filename="standup.md",
                mime_type="text/markdown",
            )

            assert result["life_category"] == "WORK"
            assert result["estimated_duration_minutes"] == 30
            assert result["recurrence_hint"] == "every Tuesday"
            assert result["location"] == "https://meet.example.com/abc"
            assert result["reminder_offset_minutes"] == 15
            assert result["urls"] == ["https://meet.example.com/abc"]
            assert result["contacts"][0]["email"] == "alex@example.com"
            assert result["document_type"] == "meeting_notes"
            assert result["action_or_info"] == "action"
            assert result["language"] == "en"

            # Technical metadata is *always* attached.
            tm = result["technical_metadata"]
            assert tm is not None
            assert tm["file_size_bytes"] == len(body)
            assert tm["line_count"] == 3

    def test_empty_file_still_returns_technical_metadata_skeleton(self):
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client

            result = extract_from_file(
                file_bytes=b"",
                filename="empty.txt",
                mime_type="text/plain",
            )
            assert result["title"] == ""
            tm = result["technical_metadata"]
            assert tm is not None
            assert tm["file_size_bytes"] == 0
            assert tm["page_count"] is None
            assert tm["word_count"] is None
            mock_client.messages.create.assert_not_called()

    @patch.dict("os.environ", {"ANTHROPIC_API_KEY": "sk-test-key"})
    def test_response_with_no_enrichment_keys_defaults_to_none(self):
        """If the model omits the new enrichment fields, the response still
        validates against the schema (all defaults to None / [])."""
        from app.schemas.ai import FileExtractionResponse
        from app.services.file_extraction import extract_from_file

        with patch("app.services.ai_productivity.anthropic") as mock_anthropic:
            mock_client = MagicMock()
            mock_anthropic.Anthropic.return_value = mock_client
            mock_client.messages.create.return_value = _make_mock_message({
                "title": "Quick Note",
                "subtasks": [],
                "tags": [],
                "detected_dates": [],
                "confidence": 0.4,
            })

            result = extract_from_file(
                file_bytes=b"quick note text",
                filename="note.txt",
                mime_type="text/plain",
            )
            # Pydantic round-trip — the existing test_returns_structured_suggestion_for_text_file
            # only exercises the router, this asserts the service-level dict
            # cleanly validates as the response schema.
            parsed = FileExtractionResponse.model_validate(result)
            assert parsed.life_category is None
            assert parsed.urls == []
            assert parsed.contacts == []
            assert parsed.technical_metadata is not None
            assert parsed.technical_metadata.file_size_bytes == len(b"quick note text")
